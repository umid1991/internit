import requests
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt

def webpage_to_word(url, filename="web_analytics.docx"):
    try:
        headers = {
            'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/120.0.0.0 Safari/537.36',
            'Accept-Language': 'ru-RU,ru;q=0.9,en-US;q=0.8,en;q=0.7'
        }
        
        # Session yaratish
        session = requests.Session()
        session.headers.update(headers)
        
        # Sahifani yuklab olish
        response = session.get(url, timeout=10)
        response.raise_for_status()
        
        # Agar redirect bo'lsa
        if response.history:
            print(f"Redirected to: {response.url}")
        
        soup = BeautifulSoup(response.text, 'html.parser')
        doc = Document()
        
        # Sarlavhani olish
        title = soup.find('h1').get_text() if soup.find('h1') else "Maqola sarlavhasi"
        doc.add_heading(title, level=0)
        
        # Kontentni izlash
        content = soup.find('article') or soup.find(class_='article') or soup.find(class_='content')
        
        if content:
            for element in content.find_all(['p', 'h2', 'h3', 'ul', 'ol', 'blockquote']):
                if element.name == 'p':
                    doc.add_paragraph(element.get_text())
                elif element.name in ['h2', 'h3']:
                    doc.add_heading(element.get_text(), level=int(element.name[1]))
                elif element.name in ['ul', 'ol']:
                    for li in element.find_all('li', recursive=False):
                        doc.add_paragraph(li.get_text(), style='ListBullet' if element.name == 'ul' else 'ListNumber')
                elif element.name == 'blockquote':
                    doc.add_paragraph(element.get_text(), style='IntenseQuote')
        
        doc.save(filename)
        print(f"Fayl muvaffaqiyatli saqlandi: {filename}")
    
    except requests.HTTPError as e:
        print(f"HTTP xatosi: {e.response.status_code} - {e.response.reason}")
    except Exception as e:
        print(f"Xatolik yuz berdi: {str(e)}")

if __name__ == "__main__":
    url = "https://practicum.yandex.ru/blog/chto-takoe-web-analytika"
    webpage_to_word(url)